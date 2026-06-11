# -*- coding: utf-8 -*-
from zipfile import ZipFile
from pathlib import Path
from docx import Document
from lxml import etree
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
MD_PATH = DOCS / "codex-deep-research-100-page-report-20260611.md"
DOCX_PATH = DOCS / "codex-deep-research-100-page-report-20260611.docx"
DATE = "2026-06-11"
TOTAL_PAGES = 100

sources = {
    "official": "官方资料：Codex CLI https://developers.openai.com/codex/cli；Codex MCP https://developers.openai.com/codex/mcp；Apply Patch https://developers.openai.com/api/docs/guides/tools-apply-patch；Tool Search https://developers.openai.com/api/docs/guides/tools-tool-search；Shell https://developers.openai.com/api/docs/guides/tools-shell；Web Search https://developers.openai.com/api/docs/guides/tools-web-search；Agents SDK https://developers.openai.com/api/docs/guides/agents",
    "turn": ".refs/openai-codex/codex-rs/core/src/session/turn.rs:141,163,236,640,692,770,892,1098,1106,1616",
    "tools": ".refs/openai-codex/codex-rs/core/src/tools/router.rs:36,52,59；registry.rs:161,185,273,294；spec_plan.rs:725,789,797；mcp_tool_exposure.rs:10,32,36,40,46；tool_search.rs:23,57,73,109",
    "exec": ".refs/openai-codex/codex-rs/core/src/tools/handlers/shell_spec.rs:20,103,217,229,267；unified_exec/exec_command.rs:72,222,236,279,291；write_stdin.rs:32,47,100；process_manager.rs:432,443,578,595,643,668,1093,1219",
    "patch_perm": ".refs/openai-codex/codex-rs/core/src/tools/handlers/apply_patch.rs:301,313,363,442,502,515；request_permissions.rs:32,55,64,78；permissions_instructions.rs:62",
    "output": ".refs/openai-codex/codex-rs/core/src/tools/context.rs:66,148,235,308,318,404,427；events.rs:54,167,316,337,395",
    "skills": ".refs/openai-codex/codex-rs/core-skills/src/loader.rs:56,83,89,161,812；core-plugins/src/manifest.rs:14,38,48,140,240；gpt_5_codex_prompt.md:1,5,11",
    "aigl": "F:/AIGril/electron/humanclaw-agent-runner.cjs:3544,3545,3561,3562,5175,5197,5223,5484,5607,5678,5798；humanclaw-tool-contracts.cjs:538,584,661,682；humanclaw-runtime.cjs:471,528,771,923；humanclaw-gateway.cjs:663,717,734,885,1635,1695",
    "gaia": "F:/AIGril/scripts/run-gaia-level1-lite.mjs:328,355,414,425,600,622；F:/AIGril/eval-results/engineering/gaia-official/*20260611*.summary.json；F:/AIGril/docs/aigl-vs-codex-gaia-failure-source-report-20260611.md",
}

group_meta = {
    "executive": ("EXEC", "总览判断", "P0", "0B3D5C"),
    "codex_arch": ("ARCH", "Codex 执行链", "P0", "1F4E79"),
    "tools": ("TOOLS", "工具暴露与发现", "P0", "275D38"),
    "exec": ("EXEC-RUNTIME", "Shell / Session", "P0", "7A4E00"),
    "patch_perm": ("SAFE-EDIT", "Patch / Permission", "P0", "8A3A3A"),
    "observability": ("EVIDENCE", "观测与证据", "P0", "4E4A8A"),
    "compaction": ("MEMORY", "上下文压缩", "P0", "5A6570"),
    "gaia_longchain": ("GAIA-LONG", "长链检索", "P0", "244C8A"),
    "gaia_web": ("GAIA-WEB", "网页证据定位", "P0", "006D77"),
    "gaia_media": ("GAIA-MEDIA", "视频 / 视觉", "P1", "9A5B00"),
    "gaia_structured": ("GAIA-DATA", "结构化来源", "P1", "5A6B2E"),
    "finalizer": ("FINAL", "Exact Answer", "P0", "8A2042"),
    "roadmap": ("ROADMAP", "落地路线", "P0", "30363D"),
}

category_stats = [
    ("长链检索", "6", "0", "多跳链路没有固化成锚点和下一跳状态，反复 web_search。"),
    ("网页证据定位", "2", "0", "页面内字段、格式和 DOM/渲染证据定位不足。"),
    ("视频/视觉", "4", "2", "字幕/网页可救两题，真实帧/棋盘结构化仍失败。"),
    ("专用数据源/结构化字段", "8", "0", "专用 adapter 覆盖和结构化结果传递不足。"),
    ("finalizer/观测捕获", "4", "0", "答案提取、证据捕获和格式化提交仍不可靠。"),
]

base_by_group = {
    "executive": (
        "Codex 的优势不在单句提示，而在 runtime 把输入、工具、观察、权限、补丁、事件和下一轮决策做成闭环。模型负责推理，harness 负责把推理变成可恢复、可审计、可验证的行动。",
        "AIGL 已经有 direct-tool、tool_search、recent_turn_items 和 finalizer 的雏形，但多处仍停在摘要、提示和事后补救层，尚未形成证据驱动的主循环。",
        "先把行为从 prompt 迁移到 runtime：工具 schema 直接暴露，工具结果以结构化 transcript 持久化，final answer 与用户可见话术分离。",
        "同一题重跑时，工具调用参数错误减少，关键证据可回放，finalizer 只能提交 evidence-backed exact answer。",
        "official; turn; output; aigl; gaia",
    ),
    "codex_arch": (
        "Codex 每轮不是简单追加聊天文本，而是重新组织 turn context、skills/plugins、tool exposure、ToolRouter、model request 和 tool result。工具执行后，结果以 response item/ToolOutput 形式进入后续上下文。",
        "AIGL 的主循环已能构造 directToolSpecs，但工具结果大量压缩进 recent_turn_items 和 lossless_tool_observations；这让模型可见的是摘要，而不是完整的、可引用的执行事件。",
        "把 agent loop 收敛为 buildTurnContext -> buildToolPlan -> callModel -> dispatchTool -> appendTranscript -> nextTurn，并让所有工具结果进入同一种 transcript item。",
        "任意一次 run 可以导出完整事件链：模型看到哪些工具、调用了哪个 schema、返回了哪个 call_id、下一轮是否读到同一 observation。",
        "turn; tools; output; skills; aigl",
    ),
    "tools": (
        "Codex 的工具不是自然语言建议，而是模型可见 spec 与 runtime handler 绑定的类型化对象。工具过多时会 deferred，通过 tool_search 搜索后再加载，避免 prompt 被全部 schema 淹没。",
        "AIGL 有 tool_search contract 和 Gateway 搜索能力，但 GAIA 失败里仍出现 tool_search 参数用 question 而不是 query，说明模型没有始终在真实 schema 上生成调用。",
        "benchmark/task 模式下让 tool_search、MCP、external__ 工具成为原生 callable tools；搜索结果必须返回可加载 spec，不只返回推荐文本。",
        "150 个工具注册时，首轮只暴露核心工具和 tool_search；命中工具在下一轮以完整 schema 可调用，错参率进入测试门槛。",
        "official; tools; aigl; gaia",
    ),
    "exec": (
        "Codex 的 shell 能力是 exec_command + write_stdin + ProcessManager，而不是一次性命令。它支持 session_id、轮询、stdin、yield_time_ms、max_output_tokens、原始 token 计数和输出截断。",
        "AIGL 有 computer.exec、session_start、pty_* 和 code.test 等多条执行路径，模型侧入口碎片化；长任务和脚本输出不总能稳定成为 answer-bearing evidence。",
        "新增 Codex 同型 exec_command/write_stdin 表面，旧 exec/PTY/process_read 降为内部后端；所有命令输出统一 ToolOutput 和 transcript。",
        "pnpm test、Python 脚本、交互式 REPL 都能返回 session_id；空 write_stdin 可轮询；最终 transcript 包含 exit_code、stdout、stderr、original_token_count。",
        "official; exec; output; aigl",
    ),
    "patch_perm": (
        "Codex 把 apply_patch 和 request_permissions 做成一等工具：补丁先验证再应用，权限不足时模型调用正式权限申请工具，而不是在自然语言里请求用户批准。",
        "AIGL 已有 apply_patch、request_permissions contract 和路径守卫，但普通任务仍可能把源码修改、脚本执行、权限阻断混在 meta JSON plan 里，缺少统一生命周期。",
        "将 source edit 默认路由到 apply_patch；将权限不足返回 permission_pending；grant 后原工具链可续跑，并记录 patch/permission 事件。",
        "越界 patch 被拒绝；hunk 不匹配失败；read-only profile 下写入会触发 request_permissions；grant 后继续执行而不是重开任务。",
        "official; patch_perm; aigl",
    ),
    "observability": (
        "Codex 将工具结果、事件、截断、遥测 preview 和上下文压缩分别处理：UI 可以丰富，模型上下文必须小而准；失败和拒绝也都是正式事件。",
        "AIGL 的 observation 目前仍过度依赖 preview/digest；ClinicalTrials 和 PPT 案例说明答案可能出现在过程里，但 finalizer 未拿到稳定证据槽。",
        "建立 EvidenceArtifact：原始 stdout、JSON path、网页字段、PDF 页码、截图帧、表格行都用 artifact id 保存，model-facing 只传摘要和引用。",
        "finalizer 输入含 artifact refs；如果答案只在截断外，finalizer 必须请求读取 artifact 而不是空提交或猜答案。",
        "output; exec; aigl; gaia",
    ),
    "compaction": (
        "Codex 在 pre-sampling 和 auto-compact 路径中按 token 状态、scope limit、window ordinal 和 prefill tokens 判断是否压缩；它不是等到上下文爆掉才处理。",
        "AIGL 已有 recent_turn_items 和 prompt_compaction 设计，但任务链仍可能把长日志、重复 progress、摘要 observation 多次塞回模型。",
        "把 70% 作为早压缩阈值，75-80% 生成 CODEX_MEMORY.md；active prompt 只保留最近事件、未解决证据和任务状态。",
        "长任务跑到 70% 时写入检查点；新线程可从检查点继续；无 API key、token、敏感原文泄漏。",
        "turn; aigl; F:/AIGril/docs/aigl-codex-context-compaction.md",
    ),
    "gaia_longchain": (
        "Codex 式长链不是一直搜索，而是把每一跳变成 artifact：锚定来源、抽取字段、保存下一跳实体、再进入下一步。shell/脚本/文件可把中间状态固化。",
        "长链 6 题 0/6，典型症状是反复 web_search、没有稳定 canonical source、没有把上一跳证据作为下一跳输入。",
        "实现 retrieval controller：Query -> CandidateSource -> Fetch -> Locate -> Extract -> NextHop -> EvidenceCheck -> Final。每一跳必须有 artifact id。",
        "6 道长链题中，每题至少输出 chain table；同一实体不得在无新证据时重复搜索三次；max_steps 前必须给出缺失 hop。",
        "official; gaia",
    ),
    "gaia_web": (
        "网页证据定位需要把页面当结构化对象处理：URL、DOM/文本、渲染、字段、附近上下文、格式差异都可能是证据，不应只靠搜索摘要。",
        "网页证据定位 2/2 失败，说明页面已找到不等于字段已定位；HTTP 403、格式化诗行、Word-of-the-Day 引文作者都需要页面内控制器。",
        "增加 page evidence locator：fetch/render/DOM text 三路并行，支持 CSS/regex/nearby quote/table extraction，失败时返回定位缺口。",
        "页面内字段题必须提交 URL + selector/path + snippet；格式题必须提交渲染或保留缩进的原文证据。",
        "official; gaia",
    ),
    "gaia_media": (
        "Codex 不会神奇看懂所有视频，但它更容易升级到字幕、帧采样、截图、脚本和专用工具；关键是 runtime 允许模型发现并调用真实媒体 primitive。",
        "视频/视觉 4 题只过 2 题：可被网页/字幕救的通过，鸟类同屏计数和棋盘结构化失败，说明缺 frame-to-structure 链路。",
        "补媒体 evidence pipeline：transcript first、frame sampling、OCR/vision labels、结构化状态生成、专用 solver/validator。",
        "视频题输出 transcript/caption/frame evidence；棋盘题输出 FEN 或等价棋盘状态；计数题输出帧编号与对象列表。",
        "official; exec; gaia",
    ),
    "gaia_structured": (
        "专用数据题的正确路径通常是 adapter 或稳定表格，而不是泛搜。Codex 的优势是能更早切换到 API、下载表格、写脚本、读取 JSON path。",
        "结构化来源 8/8 失败，ClinicalTrials 甚至调用了 external tool 仍提交空答案，说明 adapter correctness 和结果传递链都不稳。",
        "按来源建立 adapters：ClinicalTrials、Baseball Reference、Wikipedia FA、Cornell LII、Olympics、NPB、BASE；每个 adapter 返回 normalized evidence rows。",
        "每个 adapter 有 contract test、gold fixture、JSON path assertion；finalizer 不接受没有 source row/path 的答案。",
        "tools; exec; aigl; gaia",
    ),
    "finalizer": (
        "Codex 普通对话 final 与 benchmark exact answer 应分离。benchmark 模式需要 final_answer 原生工具或严格 schema，并把证据检查前置到主循环，而不是事后补救。",
        "finalizer/格式化 4/4 不得分：有空提交、整段中文解释提交、低置信错误提交和 b,e 格式问题。",
        "新增 final_answer 工具，字段含 answer、format_type、confidence、evidence_refs；缺证据时返回 repair_instruction，不允许提交自然语言长段。",
        "submitted_answer 只能来自 final_answer 工具；confidence 低或 evidence_refs 空时不得提交；格式 normalization 有单元测试。",
        "gaia; aigl",
    ),
    "roadmap": (
        "最值得复制的是执行链形状：原生工具、sessionized exec、结构化 patch、权限协议、工具搜索、ToolOutput、ToolEvent、压缩和证据管理。",
        "AIGL 当前最危险的不是没有工具，而是多条半成品链路并存：meta planner、direct tools、summary observation、finalizer、adapter 各自工作但没有闭环。",
        "按 P0/P1/P2 推进：P0 先主链和证据，P1 补媒体/adapter，P2 做优化、遥测、回归和长任务体验。",
        "每个 P0 改动都有 benchmark gate；没有 GAIA 分类回归报告，不合并 runtime 行为改动。",
        "gaia; F:/AIGril/AIGrilClaw/docs/codex-code-architecture-guide.md; F:/AIGril/AIGrilClaw/docs/aigl-codex-mainline-spec.md",
    ),
}

raw_topics = [
("封面：目标、边界与最终判断", "executive", "这份报告的目标是解释 Codex 如何把复杂任务做稳，并把结论转成 AIGL 可施工的 100 页路线图。最终判断：AIGL 不应继续堆 prompt，而应重建 transcript-driven evidence executor。"),
("调研方法：官方文档 + 本地源码 + GAIA 失败复盘", "executive", "调研采用三层证据：官方 OpenAI 文档确认产品/API 语义，本地 openai-codex 源码确认实现形状，AIGL/GAIA 日志确认当前失败点。"),
("五类 GAIA 结果总览", "executive", "2026-06-11 的 24 题分类跑法显示，AIGL 只有视频/视觉类中 2 题通过，其余长链、网页证据、结构化来源和 finalizer 类均未得分。"),
("为什么 direct-tool 后仍没有明显上涨", "executive", "direct-tool 只解决“能否看到工具按钮”的第一层问题；真正失败集中在工具之后的链路保持、证据持久化、字段抽取和 exact answer 提交。"),
("Codex 不是 prompt trick", "executive", "Codex 的可交付能力来自模型和 harness 的乘积：模型能想，runtime 让它能查、能改、能观察、能恢复、能安全地继续。"),
("如何阅读这份 100 页手册", "executive", "每页都是一张工程研究卡：先读结论，再看 Codex 机制、AIGL 差距、改造建议、验收点和证据。后续可以按页拆 issue。"),
("可信度分级：A/B/C", "executive", "A 级来自源码或官方文档，B 级来自多处源码的强推论，C 级是给 AIGL 的工程建议。实现时先采纳 A+B，再评估 C。"),
("术语边界：executor、transcript、evidence、finalizer", "executive", "本报告把 executor 定义为模型与 runtime 的闭环，把 transcript 定义为可回灌上下文，把 evidence 定义为可复核事实，把 finalizer 定义为 exact answer gate。"),
("Agent Turn Loop 主入口", "codex_arch", "Codex 的 turn loop 会在采样前构建 skills/plugins、运行采样请求、处理自动压缩、再构建 ToolRouter；这是一条有状态循环，不是一轮 JSON planner。"),
("Pre-sampling Compaction", "compaction", "Codex 在模型请求前会检查 token 状态并触发 compact，避免把上下文推到不可控边缘；这与用户要求的 70% 早压缩策略一致。"),
("Skills 与 Plugins 注入", "codex_arch", "Codex 把技能和插件作为能力包注入 turn，而不是把所有规则写死在一个系统提示里；这让能力发现、依赖和策略都可维护。"),
("Model Request Construction", "codex_arch", "Codex 的模型请求由 turn context、tool specs、history、compaction state 和 final output schema 共同决定；这给 runtime 保留了调度权。"),
("ToolRouter：模型可见 spec 与 handler 分离", "codex_arch", "ToolRouter 保存 model_visible_specs，同时把 response item 转成 runtime ToolCall；模型只能走 schema，不能绕过 handler。"),
("ToolRegistry：执行边界", "codex_arch", "ToolRegistry 把 payload 交给对应 runtime，并允许工具结果在 original 与 model_visible 之间做转换；这正是证据护栏的位置。"),
("Spec Plan：工具面按环境规划", "tools", "Codex 的 spec plan 统一加入 exec_command、write_stdin、apply_patch、request_permissions、tool_search 和 MCP runtime tools。"),
("MCP direct/deferred threshold", "tools", "Codex 使用 DIRECT_MCP_TOOL_EXPOSURE_THRESHOLD=100，并在工具多或 tool_search 开启时把 MCP tools deferred，避免首轮 prompt 膨胀。"),
("Tool Search 的 BM25 检索", "tools", "tool_search 是 runtime-native 工具，内部构建搜索引擎查 deferred tool metadata；它不是一句“你可以搜索工具”的提示。"),
("Tool Search 输出必须是可加载 spec", "tools", "Codex 的 ToolSearchOutput 会进入模型可见输出；AIGL 应返回 tool id、schema、source、provenance 和 call pattern，而不是自然语言建议。"),
("MCP Tool Boundary", "tools", "Codex 把 MCP server 的原始工具身份和模型可见 schema 分开管理；这能防止桥接层把不同工具混成自由文本。"),
("MCP Resources 与 Prompt Context", "tools", "MCP 不只是 function call，还包括 resources、templates 和上下文。AIGL 要避免只实现 tool bridge，却忽略资源读取和证据引用。"),
("Plugin Manifest 结构", "codex_arch", "Codex plugin manifest 能声明 skills、mcp servers、apps、hooks 和 interface；这为能力包治理提供了标准边界。"),
("Skill Loader 与依赖声明", "codex_arch", "Skill loader 读取 SKILL.md metadata、dependencies 和 policy；AIGL 的能力包也应能声明依赖工具缺失时的 degraded 状态。"),
("系统提示的正确位置", "codex_arch", "Codex prompt 强化工程习惯，例如优先 rg、使用 apply_patch、不回滚用户改动；但 prompt 只约束行为倾向，不替代 runtime。"),
("AGENTS.md 与仓库指令层", "codex_arch", "Codex 会把仓库局部指令纳入上下文；AIGL 也要支持项目级约束，但不能让它覆盖安全和 tool schema。"),
("Conversation Item 模型", "codex_arch", "Codex 把工具调用、非工具回复、事实和完成项作为可处理的 items；AIGL 应避免把所有历史压成一段自然语言状态。"),
("Tool Call ID 关联", "codex_arch", "每个 tool_call 与 tool_result 必须通过 call_id 关联；没有 call_id 的 evidence 很难被 finalizer 稳定引用。"),
("Transcript Continuation", "codex_arch", "真正的续跑不是重新描述进度，而是把上一轮 response item 和工具结果继续送入下一轮。AIGL 目前摘要链路仍会丢细节。"),
("Failure Observation", "observability", "失败工具结果也必须进入 transcript，因为下一轮恢复需要知道错误类型、参数、stderr 和可修复路径。"),
("exec_command schema", "exec", "Codex 暴露 cmd、workdir、tty、yield_time_ms、max_output_tokens、权限参数等字段；这比 AIGL 的泛化 computer.exec 更适合模型稳定调用。"),
("write_stdin 与空轮询", "exec", "write_stdin 允许 chars 为空时只轮询输出，这让长命令和交互命令可以持续观察，不需要重启进程。"),
("ProcessManager 生命周期", "exec", "ProcessManager 管理启动、收集输出、stdin 写入、完成状态和进程数量上限，是 Codex shell 能力的核心。"),
("yield_time_ms 与 deadline", "exec", "Codex 会 clamp yield time，并按 deadline 收集输出；模型不需要猜等待多久，也不会让命令无限阻塞。"),
("输出 token accounting", "exec", "Codex 计算 original_token_count，再按 max_output_tokens/truncation policy 返回模型可见文本；AIGL 的日志/摘要也应保留原始规模元数据。"),
("Session Cap 与资源治理", "exec", "长进程能力必须配套最大进程数和清理策略，否则 agent 会在 benchmark 或长任务中泄漏资源。"),
("Windows Sandbox 注意事项", "exec", "本项目在 Windows 上运行，Codex 源码有 Windows sandbox 分支；AIGL 的执行层也要把路径、权限、shell 语义作为平台差异处理。"),
("Shell Tool 不是裸 shell", "exec", "官方 Shell 工具语义和 Codex 源码都说明 shell 应是受控 runtime，而不是让模型拼接任意字符串后直接执行。"),
("apply_patch schema", "patch_perm", "Apply Patch 的意义是让模型提出结构化 diff，由 runtime 验证并应用；这比 echo/sed 覆盖源码安全得多。"),
("apply_patch verification", "patch_perm", "Codex apply_patch handler 会先验证参数、路径和 correctness error，再决定应用或返回错误。AIGL 应把失败暴露给模型修复。"),
("exec 中拦截 patch", "patch_perm", "Codex 即使模型把 patch 放进 shell command，也会尝试 intercept_apply_patch；这说明源码修改主路径应被 runtime 强制收口。"),
("Patch Events 与 Delta", "patch_perm", "patch success/failure 应进入 ToolEvent，并携带 changed files 或 delta；这能支撑 UI、回滚、审计和测试选择。"),
("request_permissions Tool", "patch_perm", "Codex 的 request_permissions 是正式工具，handler 会校验 permissions profile 并等待用户/系统响应。"),
("Permission Profile", "patch_perm", "权限环境应由 runtime 写进模型上下文，并由 policy 判断 allow/deny/request_permission；不应靠模型自觉避险。"),
("Approval/Sandbox Policy", "patch_perm", "安全边界包括 filesystem、network、shell、destructive command 和 prefix rule。AIGL 需要统一策略，不要让各工具各自判断。"),
("危险命令与安全命令", "patch_perm", "Codex 的 exec policy 区分 safe command、危险操作和审批路径；AIGL 至少要对删除、移动、网络、workspace 外写入做硬规则。"),
("Grant 后续跑", "patch_perm", "权限申请的价值在于 grant 后继续原任务，而不是把当前轮中断成用户手工操作。AIGL 需要 grant store 和 retry continuation。"),
("ToolOutput Family", "observability", "Codex 有 McpToolOutput、ToolSearchOutput、FunctionToolOutput、ApplyPatchToolOutput、ExecCommandToolOutput；不同工具输出语义不同。"),
("MCP Output Wrapping", "observability", "MCP 结果不应直接拍平成文本，而要保留 server/tool/content/isError 等结构；否则 finalizer 无法判断可信来源。"),
("Function Output Truncation", "observability", "Codex 对函数输出和 exec 输出使用不同 truncation policy，并保留 omission/原始规模提示。AIGL 也要区分 UI trace 与 model context。"),
("Telemetry Preview", "observability", "遥测 preview 应限制 bytes/lines；内部日志可以富，模型可见内容必须少而准。这能降低慢任务里的 model-wait 时间。"),
("Tool Events", "observability", "Begin/Success/Failure/Rejected 事件让工具执行可观察，也让失败恢复不再依赖自然语言回忆。"),
("Compaction Window", "compaction", "Codex 跟踪 auto_compact_window_ordinal 和 prefill tokens；AIGL 的长任务也应记录当前窗口和压缩前后 token 量。"),
("Inline 与 Remote Compact", "compaction", "Codex 有 inline/remote compaction 路径；AIGL 不必照搬服务，但必须有本地可恢复检查点，尤其在远程 compact 失败时。"),
("Retention Budget", "compaction", "压缩不是总结一切，而是保留目标、最新指令、未解决证据、文件路径、命令结果和下一步；这也是本线程 70% 策略的核心。"),
("Web Search Bounded Retrieval", "gaia_web", "官方 Web Search 语义是让模型获取最新信息；对 GAIA/AIGL 来说，还要把 search result、fetch、page locator 和 evidence artifact 分层。"),
("Code Mode 与本地脚本", "exec", "复杂表格/计算题应尽快切换到脚本，而不是让模型手算；Codex 的 exec/session 能把 stdout 作为 evidence。"),
("文件系统 Artifact", "observability", "下载的 PDF、HTML、CSV、PPTX、截图、字幕和脚本输出都应成为 artifact，并有 hash/path/来源元数据。"),
("EvidenceArtifact 总体设计", "observability", "EvidenceArtifact 是 AIGL 追上 Codex 稳定性的关键抽象：把 answer-bearing field 从工具摘要里解放出来。"),
("stdout 作为证据", "observability", "很多 GAIA 附件题答案来自脚本 stdout；若 stdout 只在 preview 或中文总结里出现，finalizer 就会提交错或空。"),
("PDF/Document Evidence", "observability", "PDF 证据必须带页码、section、附近上下文和抽取器版本；只说“找到巨大 PDF”不能支撑 exact answer。"),
("Image/Vision Evidence", "gaia_media", "图像证据必须从 caption 走向 structured state，例如 bounding boxes、OCR tokens、FEN、object count 和 frame timestamp。"),
("长链检索控制器总览", "gaia_longchain", "长链任务需要状态机，不需要更多泛搜提示。每一跳都应记录 source、field、entity、confidence 和 next query。"),
("长链：论文/作者历史", "gaia_longchain", "46719c30 一类任务需要先锚定 canonical paper，再抽作者历史，而不是反复搜索标题碎片。"),
("长链：实体解析器", "gaia_longchain", "305ac316 一类题应拆成语言线索、人物实体、作品角色、最终属性四步，每步都有证据。"),
("长链：链接跟随", "gaia_longchain", "840bfca7 一类题要求从文章链接到论文、致谢、NASA award；控制器应跟随 URL，而不是回到搜索页。"),
("长链：源获取优先", "gaia_longchain", "b816bfce/bda648d7 类纸面来源题，早期应转入 library/PDF/title-specific acquisition，而不是继续 broad web_search。"),
("网页字段定位", "gaia_web", "5188369a 类题说明页面找到了不代表字段找到了；locator 要能在页面内搜索 quote、author、label 和邻近节点。"),
("网页格式证据", "gaia_web", "23dd907f 类题涉及诗行缩进/格式；纯文本提取可能丢失答案，必须保留渲染或格式化 source。"),
("媒体：字幕优先", "gaia_media", "9d191bce 和 0383a3ee 能过，说明网页/字幕证据可救一部分视频题；AIGL 应把 transcript acquisition 做成标准路径。"),
("媒体：帧采样", "gaia_media", "a1e91b78 鸟类同屏计数失败说明缺 frame sampler。控制器需要按时间窗口采样、去重、计数并输出帧证据。"),
("视觉到结构化：棋盘", "gaia_media", "cca530fc 棋盘题失败说明视觉描述不够；需要 board extractor、FEN 生成、合法性校验和 chess engine/solver。"),
("结构化来源总览", "gaia_structured", "8 道结构化来源题全败，说明 AIGL 的核心缺口不是答案推理，而是稳定进入正确数据源和字段路径。"),
("ClinicalTrials Adapter", "gaia_structured", "a0068077 已出现 external__clinicaltrials__get_study，但提交空答案；adapter 需要字段 normalization，finalizer 需要看到 enrollment path。"),
("Baseball Reference Adapter", "gaia_structured", "3f57289b 需要 baseball stats 表格，不适合 loose search；adapter 应抽 roster/stat row，并在脚本里计算 ratio/field。"),
("Wikipedia Featured Article Adapter", "gaia_structured", "4fc2f1ae 需要 FA nomination/promotion history；应走 Wikipedia 页面历史、talk/FA logs 或特定结构抓取。"),
("Cornell LII Legal Adapter", "gaia_structured", "7673d772 需要法律文本 amendment diff；adapter 应定位 rule article、版本/修订说明和目标术语。"),
("Olympics Table Adapter", "gaia_structured", "cf106601 需要 1928 Olympics 参赛/代表团表；应通过稳定表格或 Wikidata/IOC source 结构化抽取。"),
("NPB Roster Adapter", "gaia_structured", "a0c07678 需要日本棒球名单与背号邻接；adapter 要处理姓名罗马字/日文、队伍、赛季和背号排序。"),
("BASE Interactive Table", "gaia_structured", "72e110e7 属于动态/交互表格；需要 browser/table extraction 或 API，而不是搜索结果里的国家名猜测。"),
("Presidents Geo Compute", "gaia_structured", "c365c1c7 需要完整列表、坐标和距离计算；正确路径是获取 source list 后脚本计算，而非模型凭常识。"),
("通用表格抽取", "gaia_structured", "所有表格题都应输出 rows、columns、source_url、row_selector、normalization 和 computed_result；这让 finalizer 不必猜。"),
("Compute Handoff", "exec", "一旦问题需要排序、计数、距离、比例或跨表 join，应进入 script artifact。Codex 的强项是自然切到 shell/脚本并读取 stdout。"),
("Benchmark Final Answer Contract", "finalizer", "GAIA 模式必须把用户可见回复和 submitted_answer 分离；自然语言解释不能直接成为提交答案。"),
("Finalizer Confidence Gate", "finalizer", "d0633230 显示低置信错误仍提交；finalizer 必须有 confidence/evidence gate，低于阈值返回 repair 而非提交。"),
("Answer Normalization", "finalizer", "6f37996b 的 b,e 格式问题说明 exact-answer normalization 要覆盖符号集合、大小写、单位、逗号空格和数字格式。"),
("Observation Digest Gap", "observability", "所谓 lossless observation 只保留有限窗口和摘要时就不是真 lossless；答案字段一旦不在窗口里，finalizer 会失明。"),
("recent_turn_items Gap", "compaction", "recent_turn_items 是好方向，但它应承载结构化 observation refs，而不是把所有证据压成 preview 字符串。"),
("Lossless 并不 Lossless", "observability", "当工具结果被截断到 5000/1600 字符摘要时，critical field 可能丢失；应把原文存 artifact，把摘要只当索引。"),
("Direct Native Tools Gap", "tools", "AIGL direct mode 已开始暴露 native direct tools，但 fallback/meta JSON 仍会让 nested args 脆弱，尤其是 MCP/external 工具。"),
("Nested Args Schema Error", "tools", "tool_search 用 question 而不是 query 的错误说明 schema 约束没在生成时生效；真实 function schema 暴露比 prompt 纠错更可靠。"),
("tool_search query/q Bug", "tools", "AIGL contract 明确要求 query/q，但模型仍错参；修复方向是让 tool_search 成为原生工具，并用 schema validation 生成可修复 observation。"),
("Adapter Correctness", "gaia_structured", "即使工具被发现，如果 adapter 参数、字段路径或返回标准错了，最终仍失败；每个 adapter 都需要 fixture 和 contract test。"),
("P0：真正 direct-tool executor", "roadmap", "第一优先级是 benchmark/task 模式去掉 meta planner 主路径，让核心工具和命中的 external/MCP 工具直接作为 callable schema 暴露。"),
("P0：EvidenceArtifact Store", "roadmap", "第二优先级是证据仓库：stdout、网页字段、PDF 页码、JSON path、表格行、视频帧都要可引用、可复读、可传给 finalizer。"),
("P0：final_answer 原生工具", "finalizer", "第三优先级是 final_answer 工具：只接受 exact answer、evidence_refs、confidence、format_type；用户可见话术另走 presenter。"),
("P0：Retrieval Controllers", "gaia_longchain", "第四优先级是为长链、网页字段和结构化来源加 controller；禁止无新证据重复泛搜，把 max_steps 用在推进链路上。"),
("P1：Media/Vision Primitives", "gaia_media", "视频/视觉要补 transcript、frame sampler、OCR、object counting、board/state extraction 和 solver 接口；否则真实视觉题仍无法稳定通过。"),
("P1：Source Adapters", "gaia_structured", "专用数据源按失败类别补 adapter，不追求一次全覆盖，先做 ClinicalTrials、Wikipedia、Baseball、Legal、Olympics 的 gold fixtures。"),
("P1：Compaction 与 Token Telemetry", "compaction", "AIGL 需要每轮记录 prompt tokens、tool output tokens、omitted items、active window 和 artifact refs；70% 触发检查点。"),
("Regression Suite Matrix", "roadmap", "每个修复必须跑分类回归：长链 6、网页 2、视频 4、结构化 8、finalizer 4；不能只看单题变好。"),
("实施顺序与 Go/No-Go", "roadmap", "Go 条件：direct schema 调用、artifact-backed finalizer、分类回归不退化。No-Go 条件：只改 prompt、只加工具名、没有证据链和测试。"),
]


def make_pages():
    pages = []
    for n, (title, group, conclusion) in enumerate(raw_topics, 1):
        codex, aigl, fix, accept, evidence = base_by_group[group]
        code, label, priority, color = group_meta[group]
        pages.append({
            "n": n,
            "title": title,
            "group": group,
            "group_code": code,
            "group_label": label,
            "priority": priority,
            "accent": color,
            "conclusion": conclusion,
            "codex": codex,
            "aigl": aigl,
            "fix": fix,
            "accept": accept,
            "evidence": evidence,
            "evidence_detail": expand_evidence(evidence),
        })
    assert len(pages) == TOTAL_PAGES
    return pages


CJK_FONT = "Microsoft YaHei"
LATIN_FONT = "Calibri"
BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(20, 33, 45)
MUTED = RGBColor(95, 105, 115)
BORDER = "D8DEE8"
FILL_BLUE = "E8EEF5"
FILL_GRAY = "F4F6F9"
FILL_WARM = "FFF8E8"


def expand_evidence(evidence):
    expanded = []
    seen = set()
    for raw_key in evidence.split(";"):
        key = raw_key.strip()
        if not key or key in seen:
            continue
        seen.add(key)
        expanded.append(f"{key}: {sources.get(key, key)}")
    return " | ".join(expanded)


def hex_to_rgb(value):
    value = value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_run_font(run, size=None, color=None, bold=None, italic=None, name=CJK_FONT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph(p, before=0, after=4, line=1.05, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align


def shade_paragraph(p, fill):
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def border_paragraph(p, color=BORDER, size="8", space="4"):
    p_pr = p._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_label_para(doc, label, text):
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=3, line=1.06)
    r = p.add_run(label + "：")
    set_run_font(r, size=8.8, color=BLUE, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=8.8, color=DARK)
    return p


def set_cell_text(cell, text, bold=False, fill=None, size=8.2, color=DARK):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    set_cell_margins(cell)
    set_cell_border(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph(p, after=0, line=1.02)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)


def add_summary_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [1.1, 0.55, 0.55, 4.3])
    headers = ["类别", "题数", "通过", "主要结论"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, fill=FILL_BLUE)
    for row in category_stats:
        cells = table.add_row().cells
        set_table_widths(table, [1.1, 0.55, 0.55, 4.3])
        for i, txt in enumerate(row):
            set_cell_text(cells[i], txt)


def add_badges(doc, p):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [1.65, 3.35, 1.5])
    labels = [
        f"{p['group_code']}",
        p["group_label"],
        f"{p['priority']} / Page {p['n']:03d}",
    ]
    fills = [p["accent"], FILL_GRAY, FILL_WARM]
    colors = [RGBColor(255, 255, 255), DARK, RGBColor(122, 90, 0)]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_text(cell, labels[idx], bold=True, fill=fills[idx], size=7.5, color=colors[idx])


def add_compare_table(doc, p):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [3.25, 3.25])
    headers = [("Codex 机制", p["codex"]), ("AIGL 差距", p["aigl"])]
    for idx, (label, text) in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, label, bold=True, fill=FILL_BLUE, size=8.0, color=BLUE)
        body = table.rows[1].cells[idx]
        set_cell_text(body, text, fill="FFFFFF", size=8.0, color=DARK)


def add_action_table(doc, p):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [3.25, 3.25])
    headers = [("改造建议", p["fix"]), ("验收点", p["accept"])]
    for idx, (label, text) in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, label, bold=True, fill=FILL_WARM, size=8.0, color=RGBColor(122, 90, 0))
        body = table.rows[1].cells[idx]
        set_cell_text(body, text, fill="FFFFFF", size=8.0, color=DARK)


def add_page(doc, p):
    kicker = doc.add_paragraph()
    set_paragraph(kicker, before=0, after=1, line=1.0)
    r = kicker.add_run(f"CODEX 深度调研 | {p['group_code']} | Page {p['n']:03d}/100 | {DATE}")
    set_run_font(r, size=7.5, color=MUTED, bold=True)

    title = doc.add_paragraph()
    set_paragraph(title, before=0, after=3, line=1.0)
    rt = title.add_run(f"第 {p['n']:03d} 页｜{p['title']}")
    set_run_font(rt, size=13.2, color=hex_to_rgb(p["accent"]), bold=True)
    border_paragraph(title, color=p["accent"], size="8", space="2")

    add_badges(doc, p)

    lead = doc.add_paragraph()
    set_paragraph(lead, before=1, after=4, line=1.08)
    shade_paragraph(lead, FILL_GRAY)
    lr = lead.add_run("核心结论：")
    set_run_font(lr, size=9.0, color=BLUE, bold=True)
    lr2 = lead.add_run(p["conclusion"])
    set_run_font(lr2, size=9.0, color=DARK)

    if p["n"] == 3:
        add_summary_table(doc)
        spacer = doc.add_paragraph()
        set_paragraph(spacer, after=2)

    add_compare_table(doc, p)
    spacer = doc.add_paragraph()
    set_paragraph(spacer, after=1)
    add_action_table(doc, p)
    ev = add_label_para(doc, "证据键", p["evidence"])
    for run in ev.runs:
        set_run_font(run, size=7.0, color=MUTED, bold=run.bold)

    note = doc.add_paragraph()
    set_paragraph(note, before=2, after=0, line=1.0)
    nr = note.add_run("施工提示：")
    set_run_font(nr, size=7.4, color=RGBColor(122, 90, 0), bold=True)
    nr2 = note.add_run("本页若进入实现阶段，先写一个能失败的验收测试，再改 runtime；不要只追加 prompt 规则。")
    set_run_font(nr2, size=7.4, color=MUTED)


def build_markdown(pages):
    lines = [
        "# CODEX 详细调研：AIGL 稳定执行链 100 页报告",
        "",
        f"生成日期：{DATE}",
        "",
        "说明：本 Markdown 与 DOCX 同源。DOCX 使用固定 100 页研究卡片结构；每页给出结论、分组、优先级、Codex 机制、AIGL 差距、改造建议、验收点和证据。",
        "",
        "## 快速导航",
        "",
        "| 页码 | 分组 | 主题 | 优先级 |",
        "|---|---|---|---|",
    ]
    for p in pages:
        lines.append(f"| {p['n']:03d} | {p['group_label']} | {p['title']} | {p['priority']} |")
    lines.extend([
        "",
        "## GAIA 分类结果",
        "",
        "| 类别 | 题数 | 通过 | 主要结论 |",
        "|---|---:|---:|---|",
    ])
    for row in category_stats:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} |")
    lines.extend([
        "",
        "## 来源键说明",
        "",
    ])
    for key, detail in sources.items():
        lines.append(f"- `{key}`：{detail}")
    lines.append("")
    lines.extend([
        "## 正文",
        "",
    ])
    for p in pages:
        lines.extend([
            "---",
            "",
            f"## 第 {p['n']:03d} 页｜{p['title']}",
            "",
            f"**分组**：{p['group_label']} / `{p['group_code']}`",
            "",
            f"**优先级**：{p['priority']}",
            "",
            f"**结论**：{p['conclusion']}",
            "",
            f"**Codex 机制**：{p['codex']}",
            "",
            f"**AIGL 差距**：{p['aigl']}",
            "",
            f"**改造建议**：{p['fix']}",
            "",
            f"**验收点**：{p['accept']}",
            "",
            f"**证据**：{p['evidence']}",
            "",
            f"**证据展开**：{p['evidence_detail']}",
            "",
        ])
    lines.extend(["---", "", "## 官方资料索引", "", "- " + sources["official"], ""])
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_docx(pages):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = CJK_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(8.8)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(footer, after=0, line=1.0)
    fr = footer.add_run("AIGL / Codex Runtime Research Manual | 100-page source-backed draft")
    set_run_font(fr, size=7.0, color=MUTED)

    for i, p in enumerate(pages):
        add_page(doc, p)
        if i != len(pages) - 1:
            doc.add_page_break()

    props = doc.core_properties
    props.title = "CODEX 详细调研：AIGL 稳定执行链 100 页报告"
    props.subject = "Codex runtime architecture, AIGL GAIA failures, evidence executor roadmap"
    props.author = "Codex"
    props.comments = "Generated from local source-backed research on 2026-06-11."
    doc.save(DOCX_PATH)


def qa_docx():
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(DOCX_PATH) as zf:
        names = set(zf.namelist())
        xml = etree.fromstring(zf.read("word/document.xml"))
    text = "\n".join(xml.xpath(".//w:t/text()", namespaces=ns))
    page_breaks = xml.xpath('.//w:br[@w:type="page"]', namespaces=ns)
    missing = [
        f"{i:03d}"
        for i in range(1, TOTAL_PAGES + 1)
        if f"Page {i:03d}/100" not in text
    ]
    required_parts = {"word/document.xml", "docProps/core.xml", "word/styles.xml"}
    return {
        "docx_size": DOCX_PATH.stat().st_size,
        "zip_parts_required": required_parts.issubset(names),
        "explicit_page_breaks": len(page_breaks),
        "structural_pages": len(page_breaks) + 1,
        "missing_page_markers": missing,
        "question_mark_ratio": round(text.count("?") / max(len(text), 1), 5),
        "text_chars": len(text),
    }


if __name__ == "__main__":
    pages = make_pages()
    build_markdown(pages)
    build_docx(pages)
    print(MD_PATH)
    print(DOCX_PATH)
    print(qa_docx())
